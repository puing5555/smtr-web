"""MD to DOCX converter for research outputs - Korean font support"""
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

def set_korean_font(run, font_name='맑은 고딕', size=Pt(10)):
    """Set Korean-compatible font on a run"""
    run.font.size = size
    run.font.name = font_name
    # Set East Asian font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

def add_styled_paragraph(doc, text, level=None, bullet=False):
    """Add paragraph with Korean font"""
    if level is not None:
        p = doc.add_heading('', level=level)
        run = p.add_run(text)
        set_korean_font(run, size=Pt(14 - level * 2))
    elif bullet:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        set_korean_font(run)
    else:
        p = doc.add_paragraph()
        # Handle **bold** markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                set_korean_font(run)
            else:
                run = p.add_run(part)
                set_korean_font(run)
    return p

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    rFonts.set(qn('w:ascii'), '맑은 고딕')
    rFonts.set(qn('w:hAnsi'), '맑은 고딕')
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.rstrip('\n')
        
        if not line.strip():
            continue
        elif line.startswith('# '):
            add_styled_paragraph(doc, line[2:], level=1)
        elif line.startswith('## '):
            add_styled_paragraph(doc, line[3:], level=2)
        elif line.startswith('### '):
            add_styled_paragraph(doc, line[4:], level=3)
        elif line.startswith('- '):
            add_styled_paragraph(doc, line[2:], bullet=True)
        elif line.startswith('---'):
            # Horizontal rule - skip
            continue
        elif line.startswith('|') and '---' in line:
            continue
        elif line.startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            add_styled_paragraph(doc, ' | '.join(cells))
        else:
            add_styled_paragraph(doc, line)
    
    doc.save(docx_path)
    print(f"Saved: {docx_path}")

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python md_to_docx.py input.md output.docx")
        sys.exit(1)
    md_to_docx(sys.argv[1], sys.argv[2])
